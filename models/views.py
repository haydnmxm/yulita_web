from django.shortcuts import get_object_or_404, render, redirect
from .forms import ModelForm
from .models import Model, Image
from django.views.generic import DetailView, UpdateView, DeleteView, TemplateView
from docx import Document
from docx.shared import Inches
from django.http import HttpResponse, JsonResponse
from django.contrib.auth.models import User, auth
from django.contrib.auth.decorators import login_required
import os

class ModelDetailView(DetailView):
    model = Model
    template_name = "model-detail.html"
    context_object_name = 'model'

    def docx(request, pk):

        def make_str(num):
            num = str(num)
            return num

        data =  Model.objects.filter(id=pk)
        model = data[0]
        images = Image.objects.filter(model=model)
        document = Document()
        for i in data:
            document.add_heading(f'{i.name}', 0)
            p = document.add_paragraph('Код: ')
            p.add_run(i.code)
            p = document.add_paragraph('Размеры: ')
            p.add_run(i.sizes)
            p = document.add_paragraph('Расход Ткани фактический: ')
            p.add_run(make_str(i.cloth_cons_fact))
            p = document.add_paragraph('Расход Ткани Коммерческий: ')
            p.add_run(make_str(i.cloth_cons_comm))
            p = document.add_paragraph('Расход Дублерина: ',)
            p.add_run(make_str(i.dublerin_cons))
            p = document.add_paragraph('Расход Флизелин: ',)
            p.add_run(make_str(i.flizelin_cons))
            p = document.add_paragraph('Ширина Ткани: ')
            p.add_run(make_str(i.cloth_width))
            p = document.add_paragraph('Ткань: ')
            p.add_run(i.cloth)
            p = document.add_paragraph('Цвета: ')
            p.add_run(i.colors)
            p = document.add_paragraph('Техническое описание: ')
            p.add_run(i.desc)
            p = document.add_paragraph('Фурнитура: ')
            p.add_run(i.furniture)
            p = document.add_paragraph('Последовательность обработки: ')
            p.add_run(i.proccesing)
            for img in range(len(images)):
                imp = Image.return_path(images[img])
                document.add_picture(f'media/{imp}', width=Inches(6), height=Inches(6))

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename={model}.docx'
        document.save(response)
        return response

class UpdateModelView(UpdateView):
    model = Model
    template_name = "model-update.html"
    form_class = ModelForm
    context_object_name = 'model'

def delete_images(request, pk):
    model = Model.objects.get(id = pk)
    item = Image.objects.filter(model=model)
    for i in item:
        os.remove(i.image.path)
    item.delete()
    return redirect('/')
    
def add_image(request, pk):
    if request.method == 'POST':
        images = request.FILES.getlist('images')
        model = Model.objects.get(id = pk)
        for image in images:
            model_image = Image.objects.create(
                model=model,
                image=image
            )
        return redirect(f'/ModelDetail/{pk}')
    model = Model.objects.get(id = pk)
    context = {
        'model' : model
    }
    return render(request,'add-image.html', context)

class DeleteModelView(DeleteView):
    model = Model
    template_name = 'model-delete.html'
    success_url = '/'

class ModelsView(TemplateView):
    template_name = 'main.html'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        objects = Model.objects.all()
        context = {
            'objects' : objects,
        } 
        return context
    
def search_view(request):
    if request.is_ajax():
        model = request.POST.get('model')
        qs = Model.objects.filter(name__icontains=model)
        res = None
        if len(qs)>0 and len(model)>0:
            data = []
            for pos in qs:
                item = {
                    'pk' : pos.pk,
                    'name' : pos.name,
                    'code' : pos.code,
                } 
                data.append(item)
            res = data
        else:
            res = 'Нет совпадений'
        return JsonResponse({'data':res})
    return JsonResponse({})

def code_search_view(request):
    if request.is_ajax():
        code = request.POST.get('code')
        qs = Model.objects.filter(code__icontains=code)
        res = None
        if len(qs)>0 and len(code)>0:
            data = []
            for pos in qs:
                item = {
                    'pk' : pos.pk,
                    'code' : pos.code,
                    'name' : pos.name,
                } 
                data.append(item)
            res = data
        else:
            res = 'Нет совпадений'
        return JsonResponse({'data':res})
    return JsonResponse({})

@login_required(login_url='login')
def home(request):
    objects = Model.objects.all()
    context = {
        'objects' : objects
    }
    return render(request, 'main.html', context)

def login(request):
    if request.method == 'POST':
        username = request.POST['username']
        password = request.POST['password']
        user = auth.authenticate(username=username,password=password)
        print(username,password,user)
        if user is not None:
            auth.login(request,user)
            return redirect('home')
        else:
            return redirect('/')

    return render(request,'login.html')

def logout(request):
    auth.logout(request)
    return render(request, 'login.html')

def add_model(request):
    if request.method == 'POST':
        form = ModelForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('home')
    form = ModelForm()
    context = {
        'form' : form
    }
    return render(request, 'add-model.html', context)

